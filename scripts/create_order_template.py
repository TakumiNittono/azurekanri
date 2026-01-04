"""
発注書テンプレートファイルを作成するスクリプト
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_order_template():
    """発注書テンプレートファイルを作成"""
    
    # テンプレートディレクトリのパス
    template_dir = Path(__file__).parent.parent / "templates" / "documents"
    template_dir.mkdir(parents=True, exist_ok=True)
    template_path = template_dir / "order_template.docx"
    
    # 新しいドキュメントを作成
    doc = Document()
    
    # セクション設定（余白）
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # タイトル
    title = doc.add_heading('発注書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    
    # 空行
    doc.add_paragraph()
    
    # 発注書番号と発注日
    header_table = doc.add_table(rows=2, cols=2)
    header_table.style = 'Table Grid'
    
    # 発注書番号
    header_table.cell(0, 0).text = '発注書番号'
    header_table.cell(0, 1).text = '{order_number}'  # プレースホルダー
    
    # 発注日
    header_table.cell(1, 0).text = '発注日'
    header_table.cell(1, 1).text = '{order_date}'  # プレースホルダー
    
    # テーブルのスタイル設定
    for row in header_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            if cell == row.cells[0]:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # 発注先情報
    doc.add_heading('発注先', level=1)
    doc.add_paragraph('業者名: {contractor_name}')
    doc.add_paragraph('住所: {contractor_address}')
    doc.add_paragraph('電話番号: {contractor_phone}')
    doc.add_paragraph('担当者名: {contractor_contact}')
    
    doc.add_paragraph()
    
    # 発注元情報
    doc.add_heading('発注元', level=1)
    doc.add_paragraph('会社名: {company_name}')
    doc.add_paragraph('住所: {company_address}')
    doc.add_paragraph('電話番号: {company_phone}')
    doc.add_paragraph('担当者名: {company_contact}')
    
    doc.add_paragraph()
    
    # 案件情報
    doc.add_heading('案件情報', level=1)
    case_table = doc.add_table(rows=5, cols=2)
    case_table.style = 'Table Grid'
    
    case_info = [
        ('案件名', '{case_name}'),
        ('修理種別', '{repair_type}'),
        ('緊急度', '{urgency}'),
        ('場所', '{location}'),
        ('貯水槽規模', '{tank_size}'),
    ]
    
    for i, (label, value) in enumerate(case_info):
        case_table.cell(i, 0).text = label
        case_table.cell(i, 1).text = value
        case_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # 発注金額
    doc.add_heading('発注金額', level=1)
    price_table = doc.add_table(rows=1, cols=2)
    price_table.style = 'Table Grid'
    price_table.cell(0, 0).text = '金額（税込）'
    price_table.cell(0, 1).text = '¥{price:,}'
    price_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    price_table.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(14)
    price_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # 納期
    doc.add_heading('納期', level=1)
    doc.add_paragraph('希望納期: {delivery_date}')
    doc.add_paragraph('作業完了予定日: {completion_date}')
    
    doc.add_paragraph()
    
    # 支払条件
    doc.add_heading('支払条件', level=1)
    doc.add_paragraph('支払方法: {payment_method}')
    doc.add_paragraph('支払期限: {payment_due_date}')
    
    doc.add_paragraph()
    
    # 備考
    doc.add_heading('備考', level=1)
    doc.add_paragraph('{notes}')
    doc.add_paragraph()
    doc.add_paragraph('※本発注書はRAGシステムによる支援情報を基に作成されています。')
    doc.add_paragraph('※最終的な金額・内容については、実際の業者との協議により決定してください。')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 承認欄
    approval_table = doc.add_table(rows=2, cols=2)
    approval_table.style = 'Table Grid'
    approval_table.cell(0, 0).text = '発注者'
    approval_table.cell(0, 1).text = '承認者'
    
    # 署名欄のスペースを確保（2行目）
    for j in range(2):
        cell = approval_table.cell(1, j)
        cell.height = Inches(1.5)
        cell.vertical_alignment = 1  # 下揃え
    
    # 保存
    doc.save(template_path)
    print(f"発注書テンプレートを作成しました: {template_path}")
    return template_path


if __name__ == "__main__":
    create_order_template()
