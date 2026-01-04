"""
見積書・発注書生成サービス
"""
from typing import Dict, Optional
from datetime import datetime
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os


class DocumentService:
    """見積書・発注書生成サービス"""
    
    def __init__(self):
        # 日本語フォントの設定（システムフォントを使用）
        # macOSの場合、ヒラギノフォントを使用
        try:
            # ヒラギノフォントのパス（macOS）
            font_paths = [
                "/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc",
                "/System/Library/Fonts/ヒラギノ角ゴシック W6.ttc",
                "/Library/Fonts/ヒラギノ角ゴシック W3.ttc",
            ]
            for font_path in font_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('Hiragino', font_path))
                    self.font_name = 'Hiragino'
                    break
            else:
                # フォントが見つからない場合はデフォルトを使用
                self.font_name = 'Helvetica'
        except Exception:
            # エラーが発生した場合はデフォルトフォントを使用
            self.font_name = 'Helvetica'
    
    def _sanitize_text(self, text: str) -> str:
        """
        テキストをサニタイズ（特殊文字をエスケープ）
        
        Args:
            text: サニタイズするテキスト
            
        Returns:
            str: サニタイズされたテキスト
        """
        if not text:
            return ""
        # HTMLエスケープ（ReportLabのParagraphはHTMLタグをサポート）
        text = text.replace("&", "&amp;")
        text = text.replace("<", "&lt;")
        text = text.replace(">", "&gt;")
        return text
    
    def generate_estimate_draft_docx(self, case_info: Dict, rag_answer: str) -> bytes:
        """
        見積書ドラフトをWord形式で生成
        
        Args:
            case_info: 案件情報
            rag_answer: RAG回答
            
        Returns:
            bytes: DOCXファイルのバイト列
        """
        try:
            doc = Document()
            
            # タイトル
            title = doc.add_heading('見積書', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 案件情報
            doc.add_heading('案件情報', level=2)
            doc.add_paragraph(f"案件名: {case_info.get('case_name', '未設定')}")
            doc.add_paragraph(f"修理種別: {case_info.get('repair_type', '未設定')}")
            doc.add_paragraph(f"緊急度: {case_info.get('urgency', '未設定')}")
            doc.add_paragraph(f"場所: {case_info.get('location', '未設定')}")
            doc.add_paragraph(f"貯水槽規模: {case_info.get('tank_size', '未設定')}")
            
            # RAG回答
            doc.add_heading('推奨業者・価格情報', level=2)
            rag_text = rag_answer[:5000] + "..." if len(rag_answer) > 5000 else rag_answer
            doc.add_paragraph(rag_text)
            
            # 備考
            doc.add_heading('備考', level=2)
            doc.add_paragraph("本見積書はRAGシステムによる支援情報を基に作成されています。")
            doc.add_paragraph("最終的な金額・内容については、実際の業者との協議により決定してください。")
            
            # 日付
            date_str = datetime.now().strftime('%Y年%m月%d日')
            doc.add_paragraph(f"作成日: {date_str}")
            
            # バイト列に変換
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            print(f"Error in generate_estimate_draft_docx: {error_detail}")
            raise
    
    def generate_order_draft_docx(self, case_info: Dict, rag_answer: str, selected_contractor: str, price: float, contractor_info: Optional[Dict] = None, notes: Optional[str] = None) -> bytes:
        """
        発注書ドラフトをWord形式で生成（テンプレートを使用）
        
        Args:
            case_info: 案件情報
            rag_answer: RAG回答
            selected_contractor: 選定した業者名
            price: 発注金額
            
        Returns:
            bytes: DOCXファイルのバイト列
        """
        try:
            # テンプレートファイルのパス
            template_path = Path(__file__).parent.parent.parent / "templates" / "documents" / "order_template.docx"
            
            # テンプレートファイルが存在する場合は使用、存在しない場合は新規作成
            if template_path.exists():
                doc = Document(str(template_path))
            else:
                doc = Document()
                # テンプレートがない場合は簡易版を作成
                title = doc.add_heading('発注書', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # プレースホルダーを置き換え
            order_date = datetime.now().strftime('%Y年%m月%d日')
            order_number = f"ORD-{datetime.now().strftime('%Y%m%d')}-{hash(case_info.get('case_name', '')) % 10000:04d}"
            
            # テンプレート内のプレースホルダーを置き換え
            replacements = {
                '{order_number}': order_number,
                '{order_date}': order_date,
                '{contractor_name}': selected_contractor,
                '{contractor_address}': '（未設定）',
                '{contractor_phone}': '（未設定）',
                '{contractor_contact}': '（未設定）',
                '{company_name}': '（未設定）',
                '{company_address}': '（未設定）',
                '{company_phone}': '（未設定）',
                '{company_contact}': '（未設定）',
                '{case_name}': case_info.get('case_name', '未設定'),
                '{repair_type}': case_info.get('repair_type', '未設定'),
                '{urgency}': case_info.get('urgency', '未設定'),
                '{location}': case_info.get('location', '未設定'),
                '{tank_size}': case_info.get('tank_size', '未設定'),
                '{price:,}': f"{price:,.0f}",
                '{delivery_date}': '（協議により決定）',
                '{completion_date}': '（協議により決定）',
                '{payment_method}': '（協議により決定）',
                '{payment_due_date}': '（協議により決定）',
                '{notes}': (notes or '') + ('\n' if notes else '') + '本発注書はRAGシステムによる支援情報を基に作成されています。\n最終的な金額・内容については、実際の業者との協議により決定してください。',
            }
            
            # ドキュメント内のすべての段落とテーブルセルを走査して置き換え
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            
            # テーブル内のセルも置き換え
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, value)
            
            # バイト列に変換
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            print(f"Error in generate_order_draft_docx: {error_detail}")
            raise
    
    def generate_estimate_draft(self, case_info: Dict, rag_answer: str) -> bytes:
        """
        見積書ドラフトを生成
        
        Args:
            case_info: 案件情報
            rag_answer: RAG回答
            
        Returns:
            bytes: PDFファイルのバイト列
        """
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            
            styles = getSampleStyleSheet()
            
            # 日本語フォントを使用するスタイルを作成
            normal_style = ParagraphStyle(
                'NormalJP',
                parent=styles['Normal'],
                fontName=self.font_name,
            )
            heading2_style = ParagraphStyle(
                'Heading2JP',
                parent=styles['Heading2'],
                fontName=self.font_name,
            )
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a237e'),
                alignment=TA_CENTER,
                spaceAfter=30,
                fontName=self.font_name,
            )
            
            # タイトル
            story.append(Paragraph("見積書", title_style))
            story.append(Spacer(1, 20))
            
            # 案件情報
            case_name = self._sanitize_text(str(case_info.get('case_name', '未設定')))
            repair_type = self._sanitize_text(str(case_info.get('repair_type', '未設定')))
            urgency = self._sanitize_text(str(case_info.get('urgency', '未設定')))
            location = self._sanitize_text(str(case_info.get('location', '未設定')))
            tank_size = self._sanitize_text(str(case_info.get('tank_size', '未設定')))
            
            story.append(Paragraph("<b>案件情報</b>", heading2_style))
            story.append(Paragraph(f"案件名: {case_name}", normal_style))
            story.append(Paragraph(f"修理種別: {repair_type}", normal_style))
            story.append(Paragraph(f"緊急度: {urgency}", normal_style))
            story.append(Paragraph(f"場所: {location}", normal_style))
            story.append(Paragraph(f"貯水槽規模: {tank_size}", normal_style))
            story.append(Spacer(1, 20))
            
            # RAG回答から推奨業者・価格情報を抽出（簡易版）
            rag_text = self._sanitize_text(rag_answer[:2000] + "..." if len(rag_answer) > 2000 else rag_answer)
            story.append(Paragraph("<b>推奨業者・価格情報</b>", heading2_style))
            story.append(Paragraph(rag_text, normal_style))
            story.append(Spacer(1, 20))
            
            # 備考
            story.append(Paragraph("<b>備考</b>", heading2_style))
            story.append(Paragraph("本見積書はRAGシステムによる支援情報を基に作成されています。", normal_style))
            story.append(Paragraph("最終的な金額・内容については、実際の業者との協議により決定してください。", normal_style))
            story.append(Spacer(1, 20))
            
            # 日付
            date_str = datetime.now().strftime('%Y年%m月%d日')
            story.append(Paragraph(f"作成日: {date_str}", normal_style))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            print(f"Error in generate_estimate_draft: {error_detail}")
            raise
    
    def generate_order_draft(self, case_info: Dict, rag_answer: str, selected_contractor: str, price: float) -> bytes:
        """
        発注書ドラフトを生成
        
        Args:
            case_info: 案件情報
            rag_answer: RAG回答
            selected_contractor: 選定した業者名
            price: 発注金額
            
        Returns:
            bytes: PDFファイルのバイト列
        """
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            
            styles = getSampleStyleSheet()
            
            # 日本語フォントを使用するスタイルを作成
            normal_style = ParagraphStyle(
                'NormalJP',
                parent=styles['Normal'],
                fontName=self.font_name,
            )
            heading2_style = ParagraphStyle(
                'Heading2JP',
                parent=styles['Heading2'],
                fontName=self.font_name,
            )
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a237e'),
                alignment=TA_CENTER,
                spaceAfter=30,
                fontName=self.font_name,
            )
            
            # タイトル
            story.append(Paragraph("発注書", title_style))
            story.append(Spacer(1, 20))
            
            # 案件情報
            case_name = self._sanitize_text(str(case_info.get('case_name', '未設定')))
            repair_type = self._sanitize_text(str(case_info.get('repair_type', '未設定')))
            urgency = self._sanitize_text(str(case_info.get('urgency', '未設定')))
            location = self._sanitize_text(str(case_info.get('location', '未設定')))
            
            story.append(Paragraph("<b>案件情報</b>", heading2_style))
            story.append(Paragraph(f"案件名: {case_name}", normal_style))
            story.append(Paragraph(f"修理種別: {repair_type}", normal_style))
            story.append(Paragraph(f"緊急度: {urgency}", normal_style))
            story.append(Paragraph(f"場所: {location}", normal_style))
            story.append(Spacer(1, 20))
            
            # 発注先
            contractor = self._sanitize_text(str(selected_contractor))
            story.append(Paragraph("<b>発注先</b>", heading2_style))
            story.append(Paragraph(f"業者名: {contractor}", normal_style))
            story.append(Spacer(1, 20))
            
            # 発注内容
            rag_text = self._sanitize_text(rag_answer[:2000] + "..." if len(rag_answer) > 2000 else rag_answer)
            story.append(Paragraph("<b>発注内容</b>", heading2_style))
            story.append(Paragraph(rag_text, normal_style))
            story.append(Spacer(1, 20))
            
            # 発注金額
            story.append(Paragraph("<b>発注金額</b>", heading2_style))
            story.append(Paragraph(f"金額: ¥{price:,.0f}", normal_style))
            story.append(Spacer(1, 20))
            
            # 備考
            story.append(Paragraph("<b>備考</b>", heading2_style))
            story.append(Paragraph("本発注書はRAGシステムによる支援情報を基に作成されています。", normal_style))
            story.append(Spacer(1, 20))
            
            # 日付
            date_str = datetime.now().strftime('%Y年%m月%d日')
            story.append(Paragraph(f"発注日: {date_str}", normal_style))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            print(f"Error in generate_order_draft: {error_detail}")
            raise


# シングルトンインスタンス
document_service = DocumentService()

